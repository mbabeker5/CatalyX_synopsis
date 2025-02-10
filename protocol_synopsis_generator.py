from typing import List, Dict, Optional, Union
from pydantic import BaseModel
import os
import json
import logging
from dotenv import load_dotenv
from openai import OpenAI
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

class StudyArm(BaseModel):
    arm_name: str
    treatment_description: str
    dosing_schedule: str

class DosageInfo(BaseModel):
    drug_name: str
    dose: str
    frequency: str
    route_of_administration: str

class ProtocolSynopsis(BaseModel):
    # Basic Study Information
    protocol_title: str
    short_title: str
    protocol_number: str
    version_and_date: str
    study_phase: str
    indication: str
    sponsor: str

    # Study Rationale
    background: str
    mechanism_of_action: str
    justification_for_study_design: str

    # Objectives & Endpoints
    primary_objectives: List[str]
    secondary_objectives: List[str]
    exploratory_objectives: Optional[List[str]]
    primary_endpoints: List[str]
    secondary_endpoints: List[str]
    exploratory_endpoints: Optional[List[str]]

    # Study Design
    design: str
    arms: List[StudyArm]
    blinding: str
    randomization: str
    study_duration: Union[str, int]

    # Population
    target_population: str
    sample_size: int
    inclusion_criteria: List[str]
    exclusion_criteria: List[str]

    # Treatments
    investigational_product: str
    comparator: Optional[str]
    dosage_and_administration: List[DosageInfo]

    # Assessments
    efficacy_assessments: List[str]
    safety_assessments: List[str]
    pharmacokinetic_assessments: Optional[List[str]]
    immunogenicity_markers: Optional[List[str]]

    # Statistical Considerations
    sample_size_justification: str
    statistical_analysis_plan: str
    interim_analysis: Optional[str]

    # Ethics & Compliance
    ethical_considerations: str
    data_monitoring: str

    # Timeline
    estimated_study_start_date: str
    estimated_study_completion_date: str
    follow_up_duration: Union[str, int]

class ProtocolSynopsisGenerator:
    def __init__(self):
        try:
            self.client = OpenAI()
            logger.info("OpenAI client initialized successfully")
        except Exception as e:
            logger.error(f"Error initializing OpenAI client: {str(e)}")
            raise

    def _read_pdf(self, pdf_path: str) -> str:
        """Read PDF and extract text content."""
        try:
            reader = PdfReader(pdf_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            logger.info(f"Successfully read PDF with {len(reader.pages)} pages")
            return text
        except Exception as e:
            logger.error(f"Error reading PDF file: {str(e)}")
            raise

    def _chunk_text(self, text: str, max_chunk_size: int = 100000) -> List[str]:
        """Split text into chunks respecting sentence boundaries.
        Using larger chunk size for GPT-4o-mini's 128k token window."""
        try:
            sentences = text.replace('\n', ' ').split('. ')
            chunks = []
            current_chunk = []
            current_size = 0
            
            for sentence in sentences:
                sentence_size = len(sentence)
                if current_size + sentence_size > max_chunk_size:
                    chunks.append('. '.join(current_chunk) + '.')
                    current_chunk = [sentence]
                    current_size = sentence_size
                else:
                    current_chunk.append(sentence)
                    current_size += sentence_size
            
            if current_chunk:
                chunks.append('. '.join(current_chunk) + '.')
            
            logger.info(f"Text split into {len(chunks)} chunks")
            return chunks
        except Exception as e:
            logger.error(f"Error chunking text: {str(e)}")
            raise

    def _process_chunk(self, chunk: str, chunk_num: int, total_chunks: int) -> Dict:
        """Process a single chunk of text using GPT-4o-mini."""
        try:
            system_prompt = """You are a clinical research expert tasked with extracting key information from a clinical trial protocol to create a synopsis.
            The input text is a section of a larger protocol document. Extract relevant information following ICH guidelines.
            Focus on maintaining technical accuracy and precision. Return the information in a JSON format exactly matching the provided schema.
            
            Required JSON Schema:
            {
                "protocol_title": "string",
                "short_title": "string",
                "protocol_number": "string",
                "version_and_date": "string",
                "study_phase": "string",
                "indication": "string",
                "sponsor": "string",
                "background": "string",
                "mechanism_of_action": "string",
                "justification_for_study_design": "string",
                "primary_objectives": ["string"],
                "secondary_objectives": ["string"],
                "exploratory_objectives": ["string"],
                "primary_endpoints": ["string"],
                "secondary_endpoints": ["string"],
                "exploratory_endpoints": ["string"],
                "design": "string",
                "arms": [
                    {
                        "arm_name": "string",
                        "treatment_description": "string",
                        "dosing_schedule": "string"
                    }
                ],
                "blinding": "string",
                "randomization": "string",
                "study_duration": "string",
                "target_population": "string",
                "sample_size": 0,
                "inclusion_criteria": ["string"],
                "exclusion_criteria": ["string"],
                "investigational_product": "string",
                "comparator": "string",
                "dosage_and_administration": [
                    {
                        "drug_name": "string",
                        "dose": "string",
                        "frequency": "string",
                        "route_of_administration": "string"
                    }
                ],
                "efficacy_assessments": ["string"],
                "safety_assessments": ["string"],
                "pharmacokinetic_assessments": ["string"],
                "immunogenicity_markers": ["string"],
                "sample_size_justification": "string",
                "statistical_analysis_plan": "string",
                "interim_analysis": "string",
                "ethical_considerations": "string",
                "data_monitoring": "string",
                "estimated_study_start_date": "string",
                "estimated_study_completion_date": "string",
                "follow_up_duration": "string"
            }
            
            IMPORTANT: You must include ALL fields in your response, even if empty.
            - Use empty string "" for missing string fields
            - Use empty array [] for missing array fields (NEVER use empty string for array fields)
            - Use 0 for missing numeric fields
            For sponsor information, extract only the name as a string, not the full details."""

            user_prompt = f"""Process part {chunk_num}/{total_chunks} of the protocol document.
            Extract all relevant information into a structured format following ICH guidelines and the exact schema provided.
            You MUST include ALL fields in the output JSON, even if empty.
            IMPORTANT: Array fields must always be arrays, even if empty (use [] not "").
            
            Text to process:
            {chunk}
            
            Return ONLY a valid JSON object matching the schema exactly. Do not include any other text or explanation."""

            response = self.client.chat.completions.create(
                model="gpt-4o-mini",  # Using GPT-4o-mini
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.1,
                max_tokens=16384,  # GPT-4o-mini's max output tokens
                response_format={ "type": "json_object" }  # Ensure JSON output
            )

            result = response.choices[0].message.content
            logger.info(f"Successfully processed chunk {chunk_num}/{total_chunks}")
            logger.debug(f"Raw response from model: {result}")
            
            try:
                parsed_result = json.loads(result)
                # Ensure all required fields are present with default values if missing
                default_structure = {
                    "protocol_title": "",
                    "short_title": "",
                    "protocol_number": "",
                    "version_and_date": "",
                    "study_phase": "",
                    "indication": "",
                    "sponsor": "",
                    "background": "",
                    "mechanism_of_action": "",
                    "justification_for_study_design": "",
                    "primary_objectives": [],
                    "secondary_objectives": [],
                    "exploratory_objectives": [],
                    "primary_endpoints": [],
                    "secondary_endpoints": [],
                    "exploratory_endpoints": [],
                    "design": "",
                    "arms": [],
                    "blinding": "",
                    "randomization": "",
                    "study_duration": "",
                    "target_population": "",
                    "sample_size": 0,
                    "inclusion_criteria": [],
                    "exclusion_criteria": [],
                    "investigational_product": "",
                    "comparator": "",
                    "dosage_and_administration": [],
                    "efficacy_assessments": [],
                    "safety_assessments": [],
                    "pharmacokinetic_assessments": [],
                    "immunogenicity_markers": [],
                    "sample_size_justification": "",
                    "statistical_analysis_plan": "",
                    "interim_analysis": "",
                    "ethical_considerations": "",
                    "data_monitoring": "",
                    "estimated_study_start_date": "",
                    "estimated_study_completion_date": "",
                    "follow_up_duration": ""
                }

                # Ensure all array fields are actually arrays
                array_fields = [
                    "primary_objectives", "secondary_objectives", "exploratory_objectives",
                    "primary_endpoints", "secondary_endpoints", "exploratory_endpoints",
                    "arms", "inclusion_criteria", "exclusion_criteria",
                    "dosage_and_administration", "efficacy_assessments", "safety_assessments",
                    "pharmacokinetic_assessments", "immunogenicity_markers"
                ]
                
                # Update default structure with parsed results
                for key, value in parsed_result.items():
                    if key in array_fields and not isinstance(value, list):
                        if value:  # If there's a value but it's not a list
                            default_structure[key] = [value]
                        else:  # If it's empty or null
                            default_structure[key] = []
                    else:
                        default_structure[key] = value
                
                return default_structure
            except json.JSONDecodeError:
                logger.error(f"Failed to parse JSON response: {result}")
                raise
                
        except Exception as e:
            logger.error(f"Error processing chunk {chunk_num}: {str(e)}")
            raise

    def _merge_info(self, base: dict, new: dict):
        """Merge new information into base dict."""
        if not base:
            # Initialize base with empty values for all fields
            base = {
                "protocol_title": "",
                "short_title": "",
                "protocol_number": "",
                "version_and_date": "",
                "study_phase": "",
                "indication": "",
                "sponsor": "",
                "background": "",
                "mechanism_of_action": "",
                "justification_for_study_design": "",
                "primary_objectives": [],
                "secondary_objectives": [],
                "exploratory_objectives": [],
                "primary_endpoints": [],
                "secondary_endpoints": [],
                "exploratory_endpoints": [],
                "design": "",
                "arms": [],
                "blinding": "",
                "randomization": "",
                "study_duration": "",
                "target_population": "",
                "sample_size": 0,
                "inclusion_criteria": [],
                "exclusion_criteria": [],
                "investigational_product": "",
                "comparator": "",
                "dosage_and_administration": [],
                "efficacy_assessments": [],
                "safety_assessments": [],
                "pharmacokinetic_assessments": [],
                "immunogenicity_markers": [],
                "sample_size_justification": "",
                "statistical_analysis_plan": "",
                "interim_analysis": "",
                "ethical_considerations": "",
                "data_monitoring": "",
                "estimated_study_start_date": "",
                "estimated_study_completion_date": "",
                "follow_up_duration": ""
            }
        
        for key, value in new.items():
            if key not in base:
                base[key] = value
            elif isinstance(value, list):
                if not isinstance(base[key], list):
                    base[key] = []
                if value:  # Only process non-empty lists
                    if isinstance(value[0], dict):  # Handle lists of dictionaries (arms, dosage_and_administration)
                        # Merge based on content rather than direct comparison
                        for new_item in value:
                            if new_item not in base[key]:  # This works because dictionaries are compared by content
                                base[key].append(new_item)
                    else:  # Handle lists of strings
                        base[key].extend(value)
                        base[key] = list(set(base[key]))  # Remove duplicates for string lists
            elif isinstance(value, str) and value and (not base[key]):
                base[key] = value
            elif isinstance(value, (int, float)) and value and (not base[key]):
                base[key] = value

    def generate_synopsis(self, pdf_path: str) -> ProtocolSynopsis:
        """Generate a protocol synopsis from a PDF protocol document."""
        try:
            logger.info(f"Starting synopsis generation for {pdf_path}")
            
            # Read PDF
            full_text = self._read_pdf(pdf_path)
            chunks = self._chunk_text(full_text)
            
            # Process each chunk
            extracted_info = None  # Start with None to ensure proper initialization
            for i, chunk in enumerate(chunks, 1):
                chunk_info = self._process_chunk(chunk, i, len(chunks))
                if extracted_info is None:
                    extracted_info = chunk_info
                else:
                    self._merge_info(extracted_info, chunk_info)

            logger.info("Successfully generated synopsis")
            logger.debug(f"Final extracted info: {json.dumps(extracted_info, indent=2)}")
            return ProtocolSynopsis(**extracted_info)
            
        except Exception as e:
            logger.error(f"Error generating synopsis: {str(e)}")
            raise

def save_to_word(synopsis: ProtocolSynopsis, output_path: str = None):
    """Save the synopsis to a formatted Word document."""
    doc = Document()
    
    # Set up document style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_heading('PROTOCOL SYNOPSIS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Basic Study Information
    doc.add_heading('Basic Study Information', 1)
    doc.add_paragraph(f'Protocol Title: {synopsis.protocol_title}')
    doc.add_paragraph(f'Short Title: {synopsis.short_title}')
    doc.add_paragraph(f'Protocol Number: {synopsis.protocol_number}')
    doc.add_paragraph(f'Version and Date: {synopsis.version_and_date}')
    doc.add_paragraph(f'Study Phase: {synopsis.study_phase}')
    doc.add_paragraph(f'Indication: {synopsis.indication}')
    doc.add_paragraph(f'Sponsor: {synopsis.sponsor}')
    
    # Study Rationale
    doc.add_heading('Study Rationale', 1)
    doc.add_paragraph('Background:', style='Heading 2')
    doc.add_paragraph(synopsis.background)
    doc.add_paragraph('Mechanism of Action:', style='Heading 2')
    doc.add_paragraph(synopsis.mechanism_of_action)
    doc.add_paragraph('Justification for Study Design:', style='Heading 2')
    doc.add_paragraph(synopsis.justification_for_study_design)
    
    # Objectives & Endpoints
    doc.add_heading('Objectives & Endpoints', 1)
    
    doc.add_paragraph('Primary Objectives:', style='Heading 2')
    for obj in synopsis.primary_objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_paragraph('Secondary Objectives:', style='Heading 2')
    for obj in synopsis.secondary_objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    if synopsis.exploratory_objectives:
        doc.add_paragraph('Exploratory Objectives:', style='Heading 2')
        for obj in synopsis.exploratory_objectives:
            doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_paragraph('Primary Endpoints:', style='Heading 2')
    for endpoint in synopsis.primary_endpoints:
        doc.add_paragraph(endpoint, style='List Bullet')
    
    doc.add_paragraph('Secondary Endpoints:', style='Heading 2')
    for endpoint in synopsis.secondary_endpoints:
        doc.add_paragraph(endpoint, style='List Bullet')
    
    if synopsis.exploratory_endpoints:
        doc.add_paragraph('Exploratory Endpoints:', style='Heading 2')
        for endpoint in synopsis.exploratory_endpoints:
            doc.add_paragraph(endpoint, style='List Bullet')
    
    # Study Design
    doc.add_heading('Study Design', 1)
    doc.add_paragraph(f'Design: {synopsis.design}')
    doc.add_paragraph(f'Blinding: {synopsis.blinding}')
    doc.add_paragraph(f'Randomization: {synopsis.randomization}')
    doc.add_paragraph(f'Study Duration: {synopsis.study_duration}')
    
    doc.add_paragraph('Study Arms:', style='Heading 2')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Arm Name'
    header_cells[1].text = 'Treatment Description'
    header_cells[2].text = 'Dosing Schedule'
    
    for arm in synopsis.arms:
        row_cells = table.add_row().cells
        row_cells[0].text = arm.arm_name
        row_cells[1].text = arm.treatment_description
        row_cells[2].text = arm.dosing_schedule
    
    # Population
    doc.add_heading('Population', 1)
    doc.add_paragraph(f'Target Population: {synopsis.target_population}')
    doc.add_paragraph(f'Sample Size: {synopsis.sample_size}')
    
    doc.add_paragraph('Inclusion Criteria:', style='Heading 2')
    for criterion in synopsis.inclusion_criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    
    doc.add_paragraph('Exclusion Criteria:', style='Heading 2')
    for criterion in synopsis.exclusion_criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    
    # Treatments
    doc.add_heading('Treatments', 1)
    doc.add_paragraph(f'Investigational Product: {synopsis.investigational_product}')
    if synopsis.comparator:
        doc.add_paragraph(f'Comparator: {synopsis.comparator}')
    
    doc.add_paragraph('Dosage and Administration:', style='Heading 2')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Drug Name'
    header_cells[1].text = 'Dose'
    header_cells[2].text = 'Frequency'
    header_cells[3].text = 'Route of Administration'
    
    for dosage in synopsis.dosage_and_administration:
        row_cells = table.add_row().cells
        row_cells[0].text = dosage.drug_name
        row_cells[1].text = dosage.dose
        row_cells[2].text = dosage.frequency
        row_cells[3].text = dosage.route_of_administration
    
    # Assessments
    doc.add_heading('Assessments', 1)
    
    doc.add_paragraph('Efficacy Assessments:', style='Heading 2')
    for assessment in synopsis.efficacy_assessments:
        doc.add_paragraph(assessment, style='List Bullet')
    
    doc.add_paragraph('Safety Assessments:', style='Heading 2')
    for assessment in synopsis.safety_assessments:
        doc.add_paragraph(assessment, style='List Bullet')
    
    if synopsis.pharmacokinetic_assessments:
        doc.add_paragraph('Pharmacokinetic Assessments:', style='Heading 2')
        for assessment in synopsis.pharmacokinetic_assessments:
            doc.add_paragraph(assessment, style='List Bullet')
    
    if synopsis.immunogenicity_markers:
        doc.add_paragraph('Immunogenicity Markers:', style='Heading 2')
        for marker in synopsis.immunogenicity_markers:
            doc.add_paragraph(marker, style='List Bullet')
    
    # Statistical Considerations
    doc.add_heading('Statistical Considerations', 1)
    doc.add_paragraph(f'Sample Size Justification: {synopsis.sample_size_justification}')
    doc.add_paragraph(f'Statistical Analysis Plan: {synopsis.statistical_analysis_plan}')
    if synopsis.interim_analysis:
        doc.add_paragraph(f'Interim Analysis: {synopsis.interim_analysis}')
    
    # Ethics & Compliance
    doc.add_heading('Ethics & Compliance', 1)
    doc.add_paragraph(f'Ethical Considerations: {synopsis.ethical_considerations}')
    doc.add_paragraph(f'Data Monitoring: {synopsis.data_monitoring}')
    
    # Timeline
    doc.add_heading('Timeline', 1)
    if synopsis.estimated_study_start_date:
        doc.add_paragraph(f'Estimated Study Start Date: {synopsis.estimated_study_start_date}')
    if synopsis.estimated_study_completion_date:
        doc.add_paragraph(f'Estimated Study Completion Date: {synopsis.estimated_study_completion_date}')
    doc.add_paragraph(f'Follow-up Duration: {synopsis.follow_up_duration}')
    
    # Save the document
    if output_path is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"protocol_synopsis_{timestamp}.docx"
    
    doc.save(output_path)
    return output_path

def main():
    import sys
    if len(sys.argv) != 2:
        print("Usage: python protocol_synopsis_generator.py <path_to_protocol.pdf>")
        sys.exit(1)

    try:
        generator = ProtocolSynopsisGenerator()
        synopsis = generator.generate_synopsis(sys.argv[1])
        
        # Save as JSON
        json_output = synopsis.model_dump_json(indent=2)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        with open(f"protocol_synopsis_{timestamp}.json", "w") as f:
            f.write(json_output)
        
        # Save as Word document
        word_path = save_to_word(synopsis)
        
        print(f"Synopsis has been saved as:")
        print(f"1. JSON: protocol_synopsis_{timestamp}.json")
        print(f"2. Word: {word_path}")
        
    except Exception as e:
        logger.error(f"Error in main: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main() 